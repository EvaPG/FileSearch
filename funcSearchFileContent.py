import os
import docx
import xlrd
from win32com import client as wc

def doSaveAsDocx(path_file,fileName,temp_path):
    word = wc.Dispatch('Word.Application')
    doc = word.Documents.Open(path_file)
    if not os.path.exists(temp_path):
        os.makedirs(temp_path)
    newFile = temp_path+'\\'+fileName+'.docx'
    doc.SaveAs(newFile, 12, False, "", True, "", False, False, False, False)
    doc.Close()
    word.Quit()
    return newFile



def txtContentFindString(path_file,keyword):
    try:
        fp = open(path_file, "r")
        strr = fp.read()
        if (strr.find(keyword) != -1):
            return True
        else:
            return False
    except Exception as err:
        return False
        pass

def wordDocxContentFindString(path_file,keyword):
    try:
        docxFile = docx.Document(path_file);
        for paragraph in docxFile.paragraphs:
            if (keyword in paragraph.text):
                return True
            else:
                continue
        for table in docxFile.tables:
            row_count = len(table.rows)
            col_count = len(table.columns)
            for i in range(0,row_count):
                for j in range(0, col_count):
                    if keyword in table.rows[i].cells[j].text:
                        return True
                    else:
                        continue
        return False
    except Exception as err:
        return False
        pass

def wordDocContentFindString(path_file,filename,keyword,temp_path):
    try:
        SaveAsDocx = doSaveAsDocx(path_file,filename,temp_path)
        docxFile = docx.Document(SaveAsDocx);
        for paragraph in docxFile.paragraphs:
            if (keyword in paragraph.text):
                return True
            else:
                continue

        for table in docxFile.tables:
            row_count = len(table.rows)
            col_count = len(table.columns)
            for i in range(0, row_count):
                for j in range(0, col_count):
                    if keyword in table.rows[i].cells[j].text:
                        return True
                    else:
                        continue
        return False
    except Exception as err:
        return False
        pass

def excelContentFindString(path_file,keyword):
    try:
        data = xlrd.open_workbook(path_file)
        for sheet in data.sheets():
            nrows = sheet.nrows
            ncols = sheet.ncols
            for i in range(0, nrows):
                rowValues = sheet.row_values(i)
                for j in range(0, ncols):
                    if (keyword in rowValues[j]):
                        return True
                    else:
                        continue
        return False
    except Exception as err:
        return False
        pass